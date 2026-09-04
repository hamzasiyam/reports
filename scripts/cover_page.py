import os
import json
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tkinter import Tk, filedialog, messagebox, Label, Entry, Button, StringVar, Toplevel, ttk
from tkcalendar import Calendar, DateEntry
from PyPDF2 import PdfMerger
from docx2pdf import convert
from datetime import datetime

PROFILE_DIR = "./profiles"
PROFILE_FILE = os.path.join(PROFILE_DIR, "cover_page_profiles.json")
OUTPUT_DIR = "./reports/processed"
SUPPLEMENTAL_DIR = os.path.join(OUTPUT_DIR, "Supplemental Files")

# Ensure the output directory exists
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

def set_cell_border(cell, **kwargs):
    """
    Set cell's border with thinner lines
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for border_name, border_attrs in kwargs.items():
        if border_attrs:
            tag = f"w:{border_name}"
            el = tcPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                tcPr.append(el)
            for key, value in border_attrs.items():
                el.set(qn(f"w:{key}"), str(value))

def generate_cover_page(title, logo_path, logo_height_percent, heading_color, body_color, shader_color, month, year, preparer_name, email, phone, note, preparation_date):
    try:
        # Format preparation date as Month Day, Year
        formatted_preparation_date = datetime.strptime(preparation_date, '%Y-%m-%d').strftime('%B %d, %Y')
        
        # Create a Word document
        doc = Document()

        # Add the logo if provided
        if logo_path:
            from PIL import Image
            logo = Image.open(logo_path)
            logo_width, logo_height = logo.size
            logo_height = int(logo_height * int(logo_height_percent) / 100)
            logo_width = int(logo_width * logo_height / logo.size[1])
            doc.add_picture(logo_path, width=Inches(logo_width / 96), height=Inches(logo_height / 96))

        # Add the title
        heading = doc.add_heading(title, 0)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor.from_string(heading_color)

        # Add the report title
        subheading = doc.add_heading('Report Information', level=1)
        run = subheading.runs[0]
        run.font.color.rgb = RGBColor.from_string(body_color)

        # Add the table with the same style as shown in the image
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        headers = ["Report Information", "Preparer Name:", "Preparer's Email", "Preparer's Office Phone", "Stats Month and Year", "Preparation Date", "Note"]
        values = ["", preparer_name, email, phone, f"{month} {year}", formatted_preparation_date, note]

        for row_idx in range(7):
            row = table.rows[row_idx]
            for col_idx in range(2):
                cell = row.cells[col_idx]
                cell.text = headers[row_idx] if col_idx == 0 else values[row_idx]
                if row_idx == 0:
                    cell.paragraphs[0].runs[0].bold = True
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), shader_color)
                    cell._element.get_or_add_tcPr().append(shading_elm)
                if row_idx != 1:
                    cell.paragraphs[0].runs[0].bold = False
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(body_color)
                set_cell_border(cell,
                    top={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                    bottom={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                    left={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                    right={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                )

        # Save the Word document
        if not os.path.exists(SUPPLEMENTAL_DIR):
            os.makedirs(SUPPLEMENTAL_DIR)

        output_path = os.path.join(SUPPLEMENTAL_DIR, f"{title}.docx")
        doc.save(output_path)

        return output_path

    except Exception as e:
        messagebox.showerror("Error", f"Error creating the cover page: {e}")

def convert_docx_to_pdf(docx_path):
    try:
        # Convert DOCX to PDF using docx2pdf
        pdf_path = docx_path.replace('.docx', '.pdf')
        convert(docx_path, pdf_path)
        return pdf_path

    except Exception as e:
        messagebox.showerror("Error", f"Error converting DOCX to PDF: {e}")

def merge_pdfs(cover_pdf_path, original_pdf_path, output_pdf_path):
    try:
        merger = PdfMerger()
        merger.append(cover_pdf_path)
        merger.append(original_pdf_path)
        merger.write(output_pdf_path)
        merger.close()
    except Exception as e:
        messagebox.showerror("Error", f"Error merging PDFs: {e}")

def save_profile(name, profile_data):
    if not os.path.exists(PROFILE_DIR):
        os.makedirs(PROFILE_DIR)

    profiles = load_profiles()
    profiles[name] = profile_data

    with open(PROFILE_FILE, "w") as f:
        json.dump(profiles, f, indent=4)

def load_profiles():
    if os.path.exists(PROFILE_FILE):
        with open(PROFILE_FILE, "r") as f:
            return json.load(f)
    return {}

def delete_profile(profile_name):
    profiles = load_profiles()
    if profile_name in profiles:
        del profiles[profile_name]
        with open(PROFILE_FILE, "w") as f:
            json.dump(profiles, f, indent=4)

def select_profile(profile_name_var, vars):
    profiles = load_profiles()
    profile_name = profile_name_var.get()
    if profile_name in profiles:
        profile = profiles[profile_name]
        vars['title'].set(profile_name)
        vars['logo_path'].set(profile["logo_path"])
        vars['logo_height_percent'].set(profile["logo_height_percent"])
        vars['heading_color'].set(profile["heading_color"])
        vars['body_color'].set(profile["body_color"])
        vars['shader_color'].set(profile["shader_color"])
        vars['month'].set(profile["month"])
        vars['year'].set(profile["year"])
        vars['preparer_name'].set(profile["preparer_name"])
        vars['email'].set(profile["email"])
        vars['phone'].set(profile["phone"])
        vars['note'].set(profile["note"])
        vars['preparation_date'].set(profile["preparation_date"])

def main():
    def browse_file(variable, filetypes):
        file_path = filedialog.askopenfilename(filetypes=filetypes)
        variable.set(file_path)

    def validate_inputs():
        if not title.get():
            messagebox.showerror("Error", "Title cannot be empty.")
            return False
        return True

    def submit():
        if not validate_inputs():
            return
        docx_path = generate_cover_page(title.get(), logo_path.get(), logo_height_percent.get(), heading_color.get(), body_color.get(), shader_color.get(), month.get(), year.get(), preparer_name.get(), email.get(), phone.get(), note.get(), preparation_date.get())
        if docx_path:
            cover_pdf_path = convert_docx_to_pdf(docx_path)
            if cover_pdf_path:
                output_pdf_path = os.path.join(OUTPUT_DIR, f"{title.get()}_with_cover.pdf")
                merge_pdfs(cover_pdf_path, original_pdf_path.get(), output_pdf_path)
                messagebox.showinfo("Success", f"PDF created successfully: {output_pdf_path}")

    def create_profile():
        profile_window = Toplevel(root)
        profile_window.title("Create Profile")
        profile_window.geometry("600x600")

        profile_name = StringVar()
        profile_logo_path = StringVar()
        profile_logo_height_percent = StringVar(value="50")
        profile_heading_color = StringVar(value="0000FF")
        profile_body_color = StringVar(value="000000")
        profile_shader_color = StringVar(value="FFD700")
        profile_month = StringVar()
        profile_year = StringVar()
        profile_preparer_name = StringVar()
        profile_email = StringVar()
        profile_phone = StringVar()
        profile_note = StringVar()
        profile_preparation_date = StringVar()

        Label(profile_window, text="Profile Name").grid(row=0, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_name, width=50).grid(row=0, column=1)

        Label(profile_window, text="Logo File").grid(row=1, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_path, width=50).grid(row=1, column=1)
        Button(profile_window, text="Browse", command=lambda: browse_file(profile_logo_path, [("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])).grid(row=1, column=2)

        Label(profile_window, text="Logo Height (%)").grid(row=2, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_height_percent, width=10).grid(row=2, column=1, sticky='w')

        Label(profile_window, text="Heading Color (hex)").grid(row=3, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_heading_color, width=10).grid(row=3, column=1, sticky='w')

        Label(profile_window, text="Body Color (hex)").grid(row=4, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_body_color, width=10).grid(row=4, column=1, sticky='w')

        Label(profile_window, text="Shader Color (hex)").grid(row=5, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_shader_color, width=10).grid(row=5, column=1, sticky='w')

        Label(profile_window, text="Month").grid(row=6, column=0, pady=10, sticky='w')
        month_combobox = ttk.Combobox(profile_window, textvariable=profile_month, values=[
            "January", "February", "March", "April", "May", "June", 
            "July", "August", "September", "October", "November", "December"], width=18)
        month_combobox.grid(row=6, column=1, sticky='w')

        Label(profile_window, text="Year").grid(row=7, column=0, pady=10, sticky='w')
        year_combobox = ttk.Combobox(profile_window, textvariable=profile_year, values=[str(y) for y in range(2022, 2051)], width=18)
        year_combobox.grid(row=7, column=1, sticky='w')

        Label(profile_window, text="Preparer Name").grid(row=8, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_preparer_name, width=50).grid(row=8, column=1)

        Label(profile_window, text="Email").grid(row=9, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_email, width=50).grid(row=9, column=1)

        Label(profile_window, text="Phone").grid(row=10, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_phone, width=50).grid(row=10, column=1)

        Label(profile_window, text="Note").grid(row=11, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_note, width=50).grid(row=11, column=1)

        Label(profile_window, text="Preparation Date").grid(row=12, column=0, pady=10, sticky='w')
        preparation_date_picker = DateEntry(profile_window, textvariable=profile_preparation_date, date_pattern='yyyy-mm-dd', width=20)
        preparation_date_picker.grid(row=12, column=1, sticky='w')

        def save_new_profile():
            profile_data = {
                "logo_path": profile_logo_path.get(),
                "logo_height_percent": profile_logo_height_percent.get(),
                "heading_color": profile_heading_color.get(),
                "body_color": profile_body_color.get(),
                "shader_color": profile_shader_color.get(),
                "month": profile_month.get(),
                "year": profile_year.get(),
                "preparer_name": profile_preparer_name.get(),
                "email": profile_email.get(),
                "phone": profile_phone.get(),
                "note": profile_note.get(),
                "preparation_date": profile_preparation_date.get()
            }
            save_profile(profile_name.get(), profile_data)
            profile_window.destroy()
            load_profile_names()

        Button(profile_window, text="Save Profile", command=save_new_profile).grid(row=13, columnspan=3, pady=20)

    def edit_profile():
        selected_profile = profile_name_var.get()
        profiles = load_profiles()
        if selected_profile not in profiles:
            messagebox.showerror("Error", "Selected profile does not exist.")
            return

        profile_window = Toplevel(root)
        profile_window.title("Edit Profile")
        profile_window.geometry("600x600")

        profile_name = StringVar(value=selected_profile)
        profile_logo_path = StringVar(value=profiles[selected_profile]["logo_path"])
        profile_logo_height_percent = StringVar(value=profiles[selected_profile]["logo_height_percent"])
        profile_heading_color = StringVar(value=profiles[selected_profile]["heading_color"])
        profile_body_color = StringVar(value=profiles[selected_profile]["body_color"])
        profile_shader_color = StringVar(value=profiles[selected_profile]["shader_color"])
        profile_month = StringVar(value=profiles[selected_profile]["month"])
        profile_year = StringVar(value=profiles[selected_profile]["year"])
        profile_preparer_name = StringVar(value=profiles[selected_profile]["preparer_name"])
        profile_email = StringVar(value=profiles[selected_profile]["email"])
        profile_phone = StringVar(value=profiles[selected_profile]["phone"])
        profile_note = StringVar(value=profiles[selected_profile]["note"])
        profile_preparation_date = StringVar(value=profiles[selected_profile]["preparation_date"])

        Label(profile_window, text="Profile Name").grid(row=0, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_name, width=50, state="readonly").grid(row=0, column=1)

        Label(profile_window, text="Logo File").grid(row=1, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_path, width=50).grid(row=1, column=1)
        Button(profile_window, text="Browse", command=lambda: browse_file(profile_logo_path, [("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])).grid(row=1, column=2)

        Label(profile_window, text="Logo Height (%)").grid(row=2, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_height_percent, width=10).grid(row=2, column=1, sticky='w')

        Label(profile_window, text="Heading Color (hex)").grid(row=3, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_heading_color, width=10).grid(row=3, column=1, sticky='w')

        Label(profile_window, text="Body Color (hex)").grid(row=4, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_body_color, width=10).grid(row=4, column=1, sticky='w')

        Label(profile_window, text="Shader Color (hex)").grid(row=5, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_shader_color, width=10).grid(row=5, column=1, sticky='w')

        Label(profile_window, text="Month").grid(row=6, column=0, pady=10, sticky='w')
        month_combobox = ttk.Combobox(profile_window, textvariable=profile_month, values=[
            "January", "February", "March", "April", "May", "June", 
            "July", "August", "September", "October", "November", "December"], width=18)
        month_combobox.grid(row=6, column=1, sticky='w')

        Label(profile_window, text="Year").grid(row=7, column=0, pady=10, sticky='w')
        year_combobox = ttk.Combobox(profile_window, textvariable=profile_year, values=[str(y) for y in range(2022, 2051)], width=18)
        year_combobox.grid(row=7, column=1, sticky='w')

        Label(profile_window, text="Preparer Name").grid(row=8, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_preparer_name, width=50).grid(row=8, column=1)

        Label(profile_window, text="Email").grid(row=9, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_email, width=50).grid(row=9, column=1)

        Label(profile_window, text="Phone").grid(row=10, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_phone, width=50).grid(row=10, column=1)

        Label(profile_window, text="Note").grid(row=11, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_note, width=50).grid(row=11, column=1)

        Label(profile_window, text="Preparation Date").grid(row=12, column=0, pady=10, sticky='w')
        preparation_date_picker = DateEntry(profile_window, textvariable=profile_preparation_date, date_pattern='yyyy-mm-dd', width=20)
        preparation_date_picker.grid(row=12, column=1, sticky='w')

        def save_edited_profile():
            profile_data = {
                "logo_path": profile_logo_path.get(),
                "logo_height_percent": profile_logo_height_percent.get(),
                "heading_color": profile_heading_color.get(),
                "body_color": profile_body_color.get(),
                "shader_color": profile_shader_color.get(),
                "month": profile_month.get(),
                "year": profile_year.get(),
                "preparer_name": profile_preparer_name.get(),
                "email": profile_email.get(),
                "phone": profile_phone.get(),
                "note": profile_note.get(),
                "preparation_date": profile_preparation_date.get()
            }
            save_profile(profile_name.get(), profile_data)
            profile_window.destroy()
            load_profile_names()

        Button(profile_window, text="Save Profile", command=save_edited_profile).grid(row=13, columnspan=3, pady=20)

    def delete_selected_profile():
        selected_profile = profile_name_var.get()
        if messagebox.askyesno("Delete Profile", f"Are you sure you want to delete the profile '{selected_profile}'?"):
            delete_profile(selected_profile)
            load_profile_names()
            clear_profile_fields()

    def clear_profile_fields():
        title.set("")
        logo_path.set("")
        logo_height_percent.set("50")
        heading_color.set("")
        body_color.set("")
        shader_color.set("")
        month.set("")
        year.set("")
        preparer_name.set("")
        email.set("")
        phone.set("")
        note.set("")
        preparation_date.set("")

    def load_profile_names():
        profiles = load_profiles()
        profile_names = list(profiles.keys())
        profile_name_combobox['values'] = profile_names

    def select_profile_callback(event):
        select_profile(profile_name_var, {
            'title': title,
            'logo_path': logo_path,
            'logo_height_percent': logo_height_percent,
            'heading_color': heading_color,
            'body_color': body_color,
            'shader_color': shader_color,
            'month': month,
            'year': year,
            'preparer_name': preparer_name,
            'email': email,
            'phone': phone,
            'note': note,
            'preparation_date': preparation_date
        })

    root = Tk()
    root.title("Cover Page Generator")
    root.geometry("600x800")

    title = StringVar()
    logo_path = StringVar()
    logo_height_percent = StringVar(value="50")
    heading_color = StringVar()
    body_color = StringVar()
    shader_color = StringVar()
    month = StringVar()
    year = StringVar()
    preparer_name = StringVar()
    email = StringVar()
    phone = StringVar()
    note = StringVar()
    preparation_date = StringVar()
    original_pdf_path = StringVar()

    Label(root, text="Title:").grid(row=0, column=0, sticky='w')
    Entry(root, textvariable=title, width=50).grid(row=0, column=1)

    Label(root, text="Logo File:").grid(row=1, column=0, sticky='w')
    Entry(root, textvariable=logo_path, width=50).grid(row=1, column=1)
    Button(root, text="Browse", command=lambda: browse_file(logo_path, [("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])).grid(row=1, column=2)

    Label(root, text="Logo Height (%):").grid(row=2, column=0, sticky='w')
    Entry(root, textvariable=logo_height_percent, width=10).grid(row=2, column=1, sticky='w')

    Label(root, text="Heading Color (hex):").grid(row=3, column=0, sticky='w')
    Entry(root, textvariable=heading_color, width=10).grid(row=3, column=1, sticky='w')

    Label(root, text="Body Text Color (hex):").grid(row=4, column=0, sticky='w')
    Entry(root, textvariable=body_color, width=10).grid(row=4, column=1, sticky='w')

    Label(root, text="Column Shading Color (hex):").grid(row=5, column=0, sticky='w')
    Entry(root, textvariable=shader_color, width=10).grid(row=5, column=1, sticky='w')

    Label(root, text="Month:").grid(row=6, column=0, sticky='w')
    month_combobox = ttk.Combobox(root, textvariable=month, values=[
        "January", "February", "March", "April", "May", "June", 
        "July", "August", "September", "October", "November", "December"], width=18)
    month_combobox.grid(row=6, column=1, sticky='w')

    Label(root, text="Year:").grid(row=7, column=0, sticky='w')
    year_combobox = ttk.Combobox(root, textvariable=year, values=[str(y) for y in range(2022, 2051)], width=18)
    year_combobox.grid(row=7, column=1, sticky='w')

    Label(root, text="Preparer Name:").grid(row=8, column=0, sticky='w')
    Entry(root, textvariable=preparer_name, width=50).grid(row=8, column=1)

    Label(root, text="Email:").grid(row=9, column=0, sticky='w')
    Entry(root, textvariable=email, width=50).grid(row=9, column=1)

    Label(root, text="Office Phone:").grid(row=10, column=0, sticky='w')
    Entry(root, textvariable=phone, width=50).grid(row=10, column=1)

    Label(root, text="Note:").grid(row=11, column=0, sticky='w')
    Entry(root, textvariable=note, width=50).grid(row=11, column=1)

    Label(root, text="Preparation Date:").grid(row=12, column=0, sticky='w')
    preparation_date_picker = DateEntry(root, textvariable=preparation_date, date_pattern='yyyy-mm-dd', width=20)
    preparation_date_picker.grid(row=12, column=1, sticky='w')

    Label(root, text="Original PDF File:").grid(row=13, column=0, sticky='w')
    Entry(root, textvariable=original_pdf_path, width=50).grid(row=13, column=1)
    Button(root, text="Browse", command=lambda: browse_file(original_pdf_path, [("PDF files", "*.pdf")])).grid(row=13, column=2)

    Button(root, text="Generate Cover Page", command=submit).grid(row=14, columnspan=3, pady=20)

    Label(root, text="Profile Name").grid(row=15, column=0, pady=10, sticky='w')
    profile_name_var = StringVar()
    profile_name_combobox = ttk.Combobox(root, textvariable=profile_name_var, state='readonly')
    profile_name_combobox.grid(row=15, column=1, pady=10, padx=10)
    profile_name_combobox.bind("<<ComboboxSelected>>", select_profile_callback)

    Button(root, text="Create Profile", command=create_profile).grid(row=15, column=2, pady=10)
    Button(root, text="Edit Profile", command=edit_profile).grid(row=16, column=1, pady=10)
    Button(root, text="Delete Profile", command=delete_selected_profile).grid(row=16, column=2, pady=10)

    load_profile_names()

    root.mainloop()

if __name__ == "__main__":
    main()
