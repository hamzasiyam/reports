import os
import json
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from tkinter import Tk, filedialog, messagebox, Label, Entry, Button, StringVar, Toplevel, ttk
import re

PROFILE_DIR = "./profiles"
PROFILE_FILE = os.path.join(PROFILE_DIR, "report_profiles.json")

# Descriptions for each sheet
descriptions = {
    "Users overview": (
        "This section details the activity of users on the website.\n\n"
        "- Unique users: The total number of distinct individuals who visited the website during the specified period. Each user is counted only once, regardless of the number of visits.\n"
        "- Sessions with new users: The number of sessions initiated by first-time visitors. Any user that has not visited the website at any time in the past is counted in this metric.\n"
        "- Sessions with returning users: The number of sessions initiated by users who have visited the website before. Any user who has visited the website at any time in the past and then returns is counted in this metric."
    ),
    "All users": (
        "Frequent users who visited the website with anonymized user ID data.\n\n"
        "- User ID: An anonymized identifier for each user.\n"
        "- No. of sessions: The number of sessions the user had on the website.\n"
        "- Country: The country from which the user accessed the website.\n"
        "- Device: The type of device the user used to access the website."
    ),
    "Insights": (
        "Insights into user behavior.\n\n"
        "- Rage clicks: Multiple clicks in quick succession, indicating frustration.\n"
        "- Dead clicks: Clicks that did not result in any action, is not inherently negative or positive.\n"
        "- Excessive scrolling: Excessive scrolling within a short period, could be an indication of difficulty finding content.\n"
        "- Quick back click: Quickly navigating back to the previous page, could be an indication of dissatisfaction.\n"
        "- No. of sessions: The number of sessions where this behavior was observed.\n"
        "- % of sessions: The percentage of total sessions where this behavior was observed."
    ),
    "Browsers": "Shows which web browsers visitors use to access the website. This information is useful for ensuring compatibility and optimizing user experience.",
    "Devices": "Details the types of devices (e.g., desktop, mobile, tablet) visitors use. This helps in tailoring the website design to different devices.",
    "Popular pages": "Identifies the most visited pages on the website.",
    "Countries": "Provides data on the geographic locations of visitors.",
    "Operating systems": "Shows which operating systems (e.g., Windows, MacOS, Android) visitors use.",
    "Smart events": "Tracks specific user interactions on the website, such as submitting a form, to provide deeper insights into user behavior.",
    "Referrer": "Details where visitors came from before arriving at the website (e.g., search engines, social media). This helps understand traffic sources.",
    "Channel": "Categorizes traffic by marketing channels (e.g., organic search, paid ads).",
    "Campaign": (
        "Tracks specific marketing campaigns driving traffic to the website. If blank, no campaign was used. "
        "A campaign can originate from various sources, including Google Ads, email marketing, social media, or any other marketing channel. "
        "To set up a campaign a tracking module ID is needed from the software used and connected to Microsoft Clarity."
    ),
    "Source": "Provides detailed information on the origin of the website traffic. This is useful for pinpointing exact traffic sources.",
    "JavaScript errors": "Logs JavaScript errors encountered by users. This helps in identifying and fixing issues affecting website functionality.",
    "Performance overview": (
        "Offers a summary of the website's performance metrics, such as load times. This helps ensure a fast and smooth user experience.\n\n"
        "- Score: An overall performance score based on various metrics.\n"
        "- Largest Contentful Paint (LCP): The time it takes for the largest content element to become visible. A key metric for user experience.\n"
        "- First Input Delay (FID): The time from when a user first interacts with the site to when the browser responds. Measures responsiveness.\n"
        "- Cumulative Layout Shift (CLS): Measures visual stability by tracking unexpected layout shifts. A low CLS ensures a smooth visual experience."
    ),
    "URL performance": (
        "Shows performance metrics for individual URLs. This helps identify and optimize slow-loading pages.\n\n"
        "- URL: The specific web address being analyzed.\n"
        "- Score: An overall performance score based on various metrics.\n"
        "- Largest Contentful Paint (LCP): The time it takes for the largest content element to become visible. A key metric for user experience.\n"
        "- First Input Delay (FID): The time from when a user first interacts with the site to when the browser responds. Measures responsiveness.\n"
        "- Cumulative Layout Shift (CLS): Measures visual stability by tracking unexpected layout shifts. A low CLS ensures a smooth visual experience."
    )
}

def set_cell_border(cell, **kwargs):
    """
    Set cell's border with thinner lines
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    for border_name, border_attrs in kwargs.items():
        if border_attrs:
            tag = f"w:{border_name}"
            el = tcPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                tcPr.append(el)
            for key, value in border_attrs.items():
                el.set(qn(f"w:{key}"), str(value))

def validate_hex_color(color):
    if re.match(r'^[0-9A-Fa-f]{6}$', color):
        return True
    return False

def generate_report(file_path, logo_path, logo_height_percent, heading_color, subheading_color, body_color, shader_color, month, year, preparer_name):
    try:
        # Read the Excel file
        xls = pd.ExcelFile(file_path)
    except Exception as e:
        messagebox.showerror("Error", f"Error reading the Excel file: {e}")
        return

    # Create a Word document
    doc = Document()

    # Add the logo if provided
    if logo_path:
        from PIL import Image
        logo = Image.open(logo_path)
        logo_width, logo_height = logo.size
        logo_height = int(logo_height * int(logo_height_percent) / 100)
        logo_width = int(logo_width * logo_height / logo.size[1])
        doc.add_picture(logo_path, width=Inches(logo_width / 96), height=Inches(logo_height / 96))

    doc.add_heading(f'Website Performance Report for {month} {year}', 0)

    # Iterate through each sheet in the Excel file
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)

        # Replace NaN values with empty strings
        df = df.replace(pd.NA, '')

        # Convert columns to string type and replace NaN with empty strings
        df = df.astype(str).fillna('')

        # Rename unnamed columns
        df.columns = [str(col) if not str(col).startswith('Unnamed') else '' for col in df.columns]

        # Add a heading for each sheet
        heading = doc.add_heading(sheet_name, level=1)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor.from_string(heading_color)

        # Add the description if available
        description = descriptions.get(sheet_name, "")
        if description:
            subheading = doc.add_paragraph(description)
            run = subheading.runs[0]
            run.font.color.rgb = RGBColor.from_string(subheading_color)

        # Add the DataFrame content to the Word document
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Table Grid'

        # Add the header row
        hdr_cells = table.rows[0].cells
        for j, column_name in enumerate(df.columns):
            hdr_cells[j].text = str(column_name)
            hdr_cells[j].paragraphs[0].runs[0].bold = True
            hdr_cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), shader_color)  # Custom shading color
            hdr_cells[j]._element.get_or_add_tcPr().append(shading_elm)

        # Add the data rows
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                cell = table.cell(i + 1, j)
                cell.text = str(df.iloc[i, j])
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(body_color)
                set_cell_border(cell,
                    top={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                    bottom={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                    left={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                    right={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                )

        # Add a page break after each sheet
        doc.add_page_break()

    # Create the output directory if it doesn't exist
    output_dir = "reports/processed"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save the Word document
    output_path = os.path.join(output_dir, f"Website Performance Insights {month} {year} Report by {preparer_name}.docx")
    doc.save(output_path)

    messagebox.showinfo("Success", f"Word document created successfully: {output_path}")

def save_profile(name, profile_data):
    if not os.path.exists(PROFILE_DIR):
        os.makedirs(PROFILE_DIR)

    profiles = load_profiles()
    profiles[name] = profile_data

    with open(PROFILE_FILE, "w") as f:
        json.dump(profiles, f, indent=4)

def load_profiles():
    if os.path.exists(PROFILE_FILE):
        with open(PROFILE_FILE, "r") as f:
            return json.load(f)
    return {}

def delete_profile(profile_name):
    profiles = load_profiles()
    if profile_name in profiles:
        del profiles[profile_name]
        with open(PROFILE_FILE, "w") as f:
            json.dump(profiles, f, indent=4)

def select_profile(profile_name_var, vars):
    profiles = load_profiles()
    profile_name = profile_name_var.get()
    if profile_name in profiles:
        profile = profiles[profile_name]
        vars['preparer_name'].set(profile_name)
        vars['logo_path'].set(profile["logo_path"])
        vars['logo_height_percent'].set(profile["logo_height_percent"])
        vars['heading_color'].set(profile["heading_color"])
        vars['subheading_color'].set(profile["subheading_color"])
        vars['body_color'].set(profile["body_color"])
        vars['shader_color'].set(profile["shader_color"])

def main():
    def browse_file(variable):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls"), ("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])
        variable.set(file_path)

    def validate_inputs():
        if not validate_hex_color(heading_color.get()):
            messagebox.showerror("Error", "Invalid color code for headings. Enter a valid 6-character hex color code.")
            return False
        if not validate_hex_color(subheading_color.get()):
            messagebox.showerror("Error", "Invalid color code for subheadings. Enter a valid 6-character hex color code.")
            return False
        if not validate_hex_color(body_color.get()):
            messagebox.showerror("Error", "Invalid color code for body text. Enter a valid 6-character hex color code.")
            return False
        if not validate_hex_color(shader_color.get()):
            messagebox.showerror("Error", "Invalid color code for column shading. Enter a valid 6-character hex color code.")
            return False
        return True

    def submit():
        if not validate_inputs():
            return
        generate_report(file_path.get(), logo_path.get(), logo_height_percent.get(), heading_color.get(), subheading_color.get(), body_color.get(), shader_color.get(), month.get(), year.get(), preparer_name.get())

    def create_profile():
        profile_window = Toplevel(root)
        profile_window.title("Create Profile")
        profile_window.geometry("400x400")

        profile_name = StringVar()
        profile_logo_path = StringVar()
        profile_logo_height_percent = StringVar(value="50")
        profile_heading_color = StringVar(value="0000FF")
        profile_subheading_color = StringVar(value="00FF00")
        profile_body_color = StringVar(value="000000")
        profile_shader_color = StringVar(value="D9D9D9")

        Label(profile_window, text="Profile Name").grid(row=0, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_name, width=50).grid(row=0, column=1)

        Label(profile_window, text="Logo File").grid(row=1, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_path, width=50).grid(row=1, column=1)
        Button(profile_window, text="Browse", command=lambda: browse_file(profile_logo_path)).grid(row=1, column=2)

        Label(profile_window, text="Logo Height (%)").grid(row=2, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_height_percent, width=10).grid(row=2, column=1, sticky='w')

        Label(profile_window, text="Heading Color (hex)").grid(row=3, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_heading_color, width=10).grid(row=3, column=1, sticky='w')

        Label(profile_window, text="Subheading Color (hex)").grid(row=4, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_subheading_color, width=10).grid(row=4, column=1, sticky='w')

        Label(profile_window, text="Body Text Color (hex)").grid(row=5, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_body_color, width=10).grid(row=5, column=1, sticky='w')

        Label(profile_window, text="Column Shading Color (hex)").grid(row=6, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_shader_color, width=10).grid(row=6, column=1, sticky='w')

        def save_new_profile():
            profile_data = {
                "logo_path": profile_logo_path.get(),
                "logo_height_percent": profile_logo_height_percent.get(),
                "heading_color": profile_heading_color.get(),
                "subheading_color": profile_subheading_color.get(),
                "body_color": profile_body_color.get(),
                "shader_color": profile_shader_color.get()
            }
            save_profile(profile_name.get(), profile_data)
            profile_window.destroy()
            load_profile_names()

        Button(profile_window, text="Save Profile", command=save_new_profile).grid(row=7, columnspan=3, pady=20)

    def edit_profile():
        selected_profile = profile_name_var.get()
        profiles = load_profiles()
        if selected_profile not in profiles:
            messagebox.showerror("Error", "Selected profile does not exist.")
            return

        profile_window = Toplevel(root)
        profile_window.title("Edit Profile")
        profile_window.geometry("600x600")

        profile_name = StringVar(value=selected_profile)
        profile_logo_path = StringVar(value=profiles[selected_profile]["logo_path"])
        profile_logo_height_percent = StringVar(value=profiles[selected_profile]["logo_height_percent"])
        profile_heading_color = StringVar(value=profiles[selected_profile]["heading_color"])
        profile_subheading_color = StringVar(value=profiles[selected_profile]["subheading_color"])
        profile_body_color = StringVar(value=profiles[selected_profile]["body_color"])
        profile_shader_color = StringVar(value=profiles[selected_profile]["shader_color"])

        Label(profile_window, text="Profile Name").grid(row=0, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_name, width=50, state="readonly").grid(row=0, column=1)

        Label(profile_window, text="Logo File").grid(row=1, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_path, width=50).grid(row=1, column=1)
        Button(profile_window, text="Browse", command=lambda: browse_file(profile_logo_path)).grid(row=1, column=2)

        Label(profile_window, text="Logo Height (%)").grid(row=2, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_height_percent, width=10).grid(row=2, column=1, sticky='w')

        Label(profile_window, text="Heading Color (hex)").grid(row=3, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_heading_color, width=10).grid(row=3, column=1, sticky='w')

        Label(profile_window, text="Subheading Color (hex)").grid(row=4, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_subheading_color, width=10).grid(row=4, column=1, sticky='w')

        Label(profile_window, text="Body Text Color (hex)").grid(row=5, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_body_color, width=10).grid(row=5, column=1, sticky='w')

        Label(profile_window, text="Column Shading Color (hex)").grid(row=6, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_shader_color, width=10).grid(row=6, column=1, sticky='w')

        def save_edited_profile():
            profile_data = {
                "logo_path": profile_logo_path.get(),
                "logo_height_percent": profile_logo_height_percent.get(),
                "heading_color": profile_heading_color.get(),
                "subheading_color": profile_subheading_color.get(),
                "body_color": profile_body_color.get(),
                "shader_color": profile_shader_color.get()
            }
            save_profile(profile_name.get(), profile_data)
            profile_window.destroy()
            load_profile_names()

        Button(profile_window, text="Save Profile", command=save_edited_profile).grid(row=7, columnspan=3, pady=20)

    def delete_selected_profile():
        selected_profile = profile_name_var.get()
        if messagebox.askyesno("Delete Profile", f"Are you sure you want to delete the profile '{selected_profile}'?"):
            delete_profile(selected_profile)
            load_profile_names()
            clear_profile_fields()

    def clear_profile_fields():
        preparer_name.set("")
        logo_path.set("")
        logo_height_percent.set("50")
        heading_color.set("")
        subheading_color.set("")
        body_color.set("")
        shader_color.set("")

    def load_profile_names():
        profiles = load_profiles()
        profile_names = list(profiles.keys())
        profile_name_combobox['values'] = profile_names

    def select_profile_callback(event):
        select_profile(profile_name_var, {
            'preparer_name': preparer_name,
            'logo_path': logo_path,
            'logo_height_percent': logo_height_percent,
            'heading_color': heading_color,
            'subheading_color': subheading_color,
            'body_color': body_color,
            'shader_color': shader_color
        })

    root = Tk()
    root.title("Website Performance Report Generator")
    root.geometry("600x800")

    file_path = StringVar()
    logo_path = StringVar()
    logo_height_percent = StringVar(value="50")
    heading_color = StringVar(value="0000FF")
    subheading_color = StringVar(value="00FF00")
    body_color = StringVar(value="000000")
    shader_color = StringVar(value="D9D9D9")
    month = StringVar()
    year = StringVar()
    preparer_name = StringVar()

    Label(root, text="Excel File:").grid(row=0, column=0, sticky='w')
    Entry(root, textvariable=file_path, width=50).grid(row=0, column=1)
    Button(root, text="Browse", command=lambda: browse_file(file_path)).grid(row=0, column=2)

    Label(root, text="Logo File:").grid(row=1, column=0, sticky='w')
    Entry(root, textvariable=logo_path, width=50).grid(row=1, column=1)
    Button(root, text="Browse", command=lambda: browse_file(logo_path)).grid(row=1, column=2)

    Label(root, text="Logo Height (%):").grid(row=2, column=0, sticky='w')
    Entry(root, textvariable=logo_height_percent, width=10).grid(row=2, column=1, sticky='w')

    Label(root, text="Heading Color (hex):").grid(row=3, column=0, sticky='w')
    Entry(root, textvariable=heading_color, width=10).grid(row=3, column=1, sticky='w')

    Label(root, text="Subheading Color (hex):").grid(row=4, column=0, sticky='w')
    Entry(root, textvariable=subheading_color, width=10).grid(row=4, column=1, sticky='w')

    Label(root, text="Body Text Color (hex):").grid(row=5, column=0, sticky='w')
    Entry(root, textvariable=body_color, width=10).grid(row=5, column=1, sticky='w')

    Label(root, text="Column Shading Color (hex):").grid(row=6, column=0, sticky='w')
    Entry(root, textvariable=shader_color, width=10).grid(row=6, column=1, sticky='w')

    Label(root, text="Month:").grid(row=7, column=0, sticky='w')
    Entry(root, textvariable=month, width=20).grid(row=7, column=1, sticky='w')

    Label(root, text="Year:").grid(row=8, column=0, sticky='w')
    Entry(root, textvariable=year, width=20).grid(row=8, column=1, sticky='w')

    Label(root, text="Preparer Name:").grid(row=9, column=0, sticky='w')
    Entry(root, textvariable=preparer_name, width=50).grid(row=9, column=1)

    Button(root, text="Generate Report", command=submit).grid(row=10, columnspan=3, pady=20)

    Label(root, text="Profile Name").grid(row=11, column=0, pady=10, sticky='w')
    profile_name_var = StringVar()
    profile_name_combobox = ttk.Combobox(root, textvariable=profile_name_var, state='readonly')
    profile_name_combobox.grid(row=11, column=1, pady=10, padx=10)
    profile_name_combobox.bind("<<ComboboxSelected>>", select_profile_callback)

    Button(root, text="Create Profile", command=create_profile).grid(row=11, column=2, pady=10)
    Button(root, text="Edit Profile", command=edit_profile).grid(row=12, column=1, pady=10)
    Button(root, text="Delete Profile", command=delete_selected_profile).grid(row=12, column=2, pady=10)

    load_profile_names()

    root.mainloop()

if __name__ == "__main__":
    main()
