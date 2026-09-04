import os
import json
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tkinter import Tk, filedialog, messagebox, Label, Entry, Button, StringVar, Toplevel, ttk
from tkcalendar import DateEntry
from PIL import ImageGrab, ImageTk, Image
from datetime import datetime

PROFILE_DIR = "./profiles"
PROFILE_FILE = os.path.join(PROFILE_DIR, "worklog_profiles.json")
OUTPUT_DIR = "./reports/processed/Work Log"
ACTIONS_FILE = "./profiles/actions.json"  # JSON file to store actions

# Ensure the profile and output directories exist
if not os.path.exists(PROFILE_DIR):
    os.makedirs(PROFILE_DIR)

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

def set_cell_border(cell, **kwargs):
    """
    Set cell's border with thinner lines
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for border_name, border_attrs in kwargs.items():
        if border_attrs:
            tag = f"w:{border_name}"
            el = tcPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                tcPr.append(el)
            for key, value in border_attrs.items():
                el.set(qn(f"w:{key}"), str(value))

def generate_report(title, action, logo_path, logo_height_percent, heading_color, body_color, shader_color, preparer_name, email, phone, note, date_of_work, clipboard_image_paths=None):
    try:
        formatted_date_of_work = datetime.strptime(date_of_work, '%Y-%m-%d').strftime('%B %d, %Y')
        month_year_folder = datetime.strptime(date_of_work, '%Y-%m-%d').strftime('%B_%Y')
        report_folder = os.path.join(OUTPUT_DIR, month_year_folder)

        # Ensure the output directory exists
        if not os.path.exists(report_folder):
            os.makedirs(report_folder)

        # Create a Word document
        doc = Document()

        # Add the logo if provided
        if logo_path:
            from PIL import Image
            logo = Image.open(logo_path)
            logo_width, logo_height = logo.size
            logo_height = int(logo_height * int(logo_height_percent) / 100)
            logo_width = int(logo_width * logo_height / logo.size[1])
            doc.add_picture(logo_path, width=Inches(logo_width / 96), height=Inches(logo_height / 96))

        # Add the title
        heading = doc.add_heading(f"{title} - {action}", 0)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor.from_string(heading_color)

        # Add the report title
        subheading = doc.add_heading('Report Information', level=1)
        run = subheading.runs[0]
        run.font.color.rgb = RGBColor.from_string(body_color)

        # Add the table
        table = doc.add_table(rows=6 if note else 5, cols=2)
        table.style = 'Table Grid'
        headers = ["Report Information", "Preparer Name:", "Preparer's Email", "Preparer's Office Phone", "Date of Work", "Note"]
        values = ["", preparer_name, email, phone, formatted_date_of_work, note]

        for row_idx in range(6 if note else 5):
            row = table.rows[row_idx]
            for col_idx in range(2):
                cell = row.cells[col_idx]
                cell.text = headers[row_idx] if col_idx == 0 else values[row_idx]
                if row_idx == 0:
                    cell.paragraphs[0].runs[0].bold = True
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), shader_color)
                    cell._element.get_or_add_tcPr().append(shading_elm)
                if row_idx != 1:
                    cell.paragraphs[0].runs[0].bold = False
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(body_color)
                set_cell_border(cell,
                                top={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                                bottom={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                                left={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                                right={"sz": 6, "val": "single", "color": "000000", "space": "0"},
                                )

        # Add the clipboard images if they exist
        if clipboard_image_paths:
            screenshots_heading = doc.add_heading('Screenshots:', level=1)
            screenshots_heading.alignment = 0  # Left align the heading
            for image_path in clipboard_image_paths:
                if image_path:
                    doc.add_picture(image_path, width=Inches(5))

        # Save the Word document
        output_path = os.path.join(report_folder, f"{title}_{action}.docx")
        doc.save(output_path)

        return output_path

    except Exception as e:
        messagebox.showerror("Error", f"Error creating the report: {e}")

def paste_image_from_clipboard(event, entry, image_labels, image_paths, index):
    try:
        image = ImageGrab.grabclipboard()
        if image:
            # Determine the correct folder for the current month and year
            month_year_folder = datetime.today().strftime('%B_%Y')
            image_folder = os.path.join(OUTPUT_DIR, month_year_folder, "Clipboard Images")

            # Ensure the "Clipboard Images" folder exists
            if not os.path.exists(image_folder):
                os.makedirs(image_folder)

            # Save the image in the "Clipboard Images" folder
            image_path = os.path.join(image_folder, f"clipboard_image_{index + 1}.png")
            image.save(image_path, "PNG")
            image_paths[index] = image_path

            # Display the image in the corresponding box
            image.thumbnail((150, 150))  # Resize the image to fit in the box
            img = ImageTk.PhotoImage(image)
            image_labels[index].config(image=img)
            image_labels[index].image = img  # Keep a reference to avoid garbage collection

            # Clear the entry text
            entry.delete(0, 'end')
        else:
            messagebox.showerror("Error", "No image in clipboard.")
    except Exception as e:
        messagebox.showerror("Error", f"Error pasting image from clipboard: {e}")

def clear_image_box(image_labels, image_paths, index):
    image_labels[index].config(image=None)
    image_labels[index].image = None
    image_paths[index] = None

def clear_all_images(image_labels, image_paths):
    for i in range(len(image_labels)):
        clear_image_box(image_labels, image_paths, i)

def save_profile(name, profile_data):
    profiles = load_profiles()
    profiles[name] = profile_data
    with open(PROFILE_FILE, "w") as f:
        json.dump(profiles, f, indent=4)

def load_profiles():
    if os.path.exists(PROFILE_FILE):
        with open(PROFILE_FILE, "r") as f:
            return json.load(f)
    return {}

def delete_profile(profile_name):
    profiles = load_profiles()
    if profile_name in profiles:
        del profiles[profile_name]
        with open(PROFILE_FILE, "w") as f:
            json.dump(profiles, f, indent=4)

def select_profile(profile_name_var, vars):
    profiles = load_profiles()
    profile_name = profile_name_var.get()
    if profile_name in profiles:
        profile = profiles[profile_name]
        vars['title'].set(profile_name)
        vars['logo_path'].set(profile["logo_path"])
        vars['logo_height_percent'].set(profile["logo_height_percent"])
        vars['heading_color'].set(profile["heading_color"])
        vars['body_color'].set(profile["body_color"])
        vars['shader_color'].set(profile["shader_color"])
        vars['preparer_name'].set(profile["preparer_name"])
        vars['email'].set(profile["email"])
        vars['phone'].set(profile["phone"])
        vars['note'].set(profile["note"])
        vars['action'].set(profile["action"])

def load_actions():
    if os.path.exists(ACTIONS_FILE):
        with open(ACTIONS_FILE, "r") as f:
            return json.load(f)
    return []

def main():
    def browse_file(variable, filetypes):
        file_path = filedialog.askopenfilename(filetypes=filetypes)
        variable.set(file_path)

    def validate_inputs():
        if not title.get():
            messagebox.showerror("Error", "Title cannot be empty.")
            return False
        return True

    def submit():
        if not validate_inputs():
            return
        report_path = generate_report(
            title.get(), action.get(), logo_path.get(), logo_height_percent.get(),
            heading_color.get(), body_color.get(), shader_color.get(),
            preparer_name.get(), email.get(), phone.get(), note.get(), date_of_work.get(), image_paths)
        if report_path:
            messagebox.showinfo("Success", f"Report created successfully: {report_path}")

    def create_profile():
        profile_window = Toplevel(root)
        profile_window.title("Create Profile")
        profile_window.geometry("1000x700")

        profile_name = StringVar()
        profile_logo_path = StringVar()
        profile_logo_height_percent = StringVar(value="50")
        profile_heading_color = StringVar(value="0000FF")
        profile_body_color = StringVar(value="000000")
        profile_shader_color = StringVar(value="FFD700")
        profile_preparer_name = StringVar()
        profile_email = StringVar()
        profile_phone = StringVar()
        profile_note = StringVar()
        profile_action = StringVar(value="Manage")

        Label(profile_window, text="Profile Name", font=("Arial", 14)).grid(row=0, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_name, width=50, font=("Arial", 14)).grid(row=0, column=1)

        Label(profile_window, text="Logo File", font=("Arial", 14)).grid(row=1, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_path, width=50, font=("Arial", 14)).grid(row=1, column=1)
        Button(profile_window, text="Browse", font=("Arial", 14), command=lambda: browse_file(profile_logo_path, [("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])).grid(row=1, column=2)

        Label(profile_window, text="Logo Height (%)", font=("Arial", 14)).grid(row=2, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_height_percent, width=10, font=("Arial", 14)).grid(row=2, column=1, sticky='w')

        Label(profile_window, text="Heading Color (hex)", font=("Arial", 14)).grid(row=3, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_heading_color, width=10, font=("Arial", 14)).grid(row=3, column=1, sticky='w')

        Label(profile_window, text="Body Color (hex)", font=("Arial", 14)).grid(row=4, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_body_color, width=10, font=("Arial", 14)).grid(row=4, column=1, sticky='w')

        Label(profile_window, text="Shader Color (hex)", font=("Arial", 14)).grid(row=5, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_shader_color, width=10, font=("Arial", 14)).grid(row=5, column=1, sticky='w')

        Label(profile_window, text="Preparer Name", font=("Arial", 14)).grid(row=6, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_preparer_name, width=50, font=("Arial", 14)).grid(row=6, column=1)

        Label(profile_window, text="Email", font=("Arial", 14)).grid(row=7, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_email, width=50, font=("Arial", 14)).grid(row=7, column=1)

        Label(profile_window, text="Phone", font=("Arial", 14)).grid(row=8, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_phone, width=50, font=("Arial", 14)).grid(row=8, column=1)

        Label(profile_window, text="Note", font=("Arial", 14)).grid(row=9, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_note, width=50, font=("Arial", 14)).grid(row=9, column=1)

        Label(profile_window, text="Action:", font=("Arial", 14)).grid(row=10, column=0, pady=10, sticky='w')
        action_combobox = ttk.Combobox(profile_window, textvariable=profile_action, values=load_actions(), width=18, font=("Arial", 14))
        action_combobox.grid(row=10, column=1, sticky='w')

        def save_new_profile():
            profile_data = {
                "logo_path": profile_logo_path.get(),
                "logo_height_percent": profile_logo_height_percent.get(),
                "heading_color": profile_heading_color.get(),
                "body_color": profile_body_color.get(),
                "shader_color": profile_shader_color.get(),
                "preparer_name": profile_preparer_name.get(),
                "email": profile_email.get(),
                "phone": profile_phone.get(),
                "note": profile_note.get(),
                "action": profile_action.get()
            }
            save_profile(profile_name.get(), profile_data)
            profile_window.destroy()
            load_profile_names()

        # Aligning the buttons in the same row
        Button(profile_window, text="Create Profile", font=("Arial", 14), command=save_new_profile).grid(row=11, column=0, pady=20, padx=10, sticky='e')
        Button(profile_window, text="Cancel", font=("Arial", 14), command=profile_window.destroy).grid(row=11, column=1, pady=20, padx=10, sticky='w')

    def edit_profile():
        selected_profile = profile_name_var.get()
        profiles = load_profiles()
        if selected_profile not in profiles:
            messagebox.showerror("Error", "Selected profile does not exist.")
            return

        profile_window = Toplevel(root)
        profile_window.title("Edit Profile")
        profile_window.geometry("850x600")

        profile_name = StringVar(value=selected_profile)
        profile_logo_path = StringVar(value=profiles[selected_profile]["logo_path"])
        profile_logo_height_percent = StringVar(value=profiles[selected_profile]["logo_height_percent"])
        profile_heading_color = StringVar(value=profiles[selected_profile]["heading_color"])
        profile_body_color = StringVar(value=profiles[selected_profile]["body_color"])
        profile_shader_color = StringVar(value=profiles[selected_profile]["shader_color"])
        profile_preparer_name = StringVar(value=profiles[selected_profile]["preparer_name"])
        profile_email = StringVar(value=profiles[selected_profile]["email"])
        profile_phone = StringVar(value=profiles[selected_profile]["phone"])
        profile_note = StringVar(value=profiles[selected_profile]["note"])
        profile_action = StringVar(value=profiles[selected_profile]["action"])

        Label(profile_window, text="Profile Name", font=("Arial", 14)).grid(row=0, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_name, width=50, font=("Arial", 14), state="readonly").grid(row=0, column=1)

        Label(profile_window, text="Logo File", font=("Arial", 14)).grid(row=1, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_path, width=50, font=("Arial", 14)).grid(row=1, column=1)
        Button(profile_window, text="Browse", font=("Arial", 14), command=lambda: browse_file(profile_logo_path, [("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])).grid(row=1, column=2)

        Label(profile_window, text="Logo Height (%)", font=("Arial", 14)).grid(row=2, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_logo_height_percent, width=10, font=("Arial", 14)).grid(row=2, column=1, sticky='w')

        Label(profile_window, text="Heading Color (hex)", font=("Arial", 14)).grid(row=3, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_heading_color, width=10, font=("Arial", 14)).grid(row=3, column=1, sticky='w')

        Label(profile_window, text="Body Color (hex)", font=("Arial", 14)).grid(row=4, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_body_color, width=10, font=("Arial", 14)).grid(row=4, column=1, sticky='w')

        Label(profile_window, text="Shader Color (hex)", font=("Arial", 14)).grid(row=5, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_shader_color, width=10, font=("Arial", 14)).grid(row=5, column=1, sticky='w')

        Label(profile_window, text="Preparer Name", font=("Arial", 14)).grid(row=6, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_preparer_name, width=50, font=("Arial", 14)).grid(row=6, column=1)

        Label(profile_window, text="Email", font=("Arial", 14)).grid(row=7, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_email, width=50, font=("Arial", 14)).grid(row=7, column=1)

        Label(profile_window, text="Phone", font=("Arial", 14)).grid(row=8, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_phone, width=50, font=("Arial", 14)).grid(row=8, column=1)

        Label(profile_window, text="Note", font=("Arial", 14)).grid(row=9, column=0, pady=10, sticky='w')
        Entry(profile_window, textvariable=profile_note, width=50, font=("Arial", 14)).grid(row=9, column=1)

        Label(profile_window, text="Action:", font=("Arial", 14)).grid(row=10, column=0, pady=10, sticky='w')
        action_combobox = ttk.Combobox(profile_window, textvariable=profile_action, values=load_actions(), width=18, font=("Arial", 14))
        action_combobox.grid(row=10, column=1, sticky='w')

        def save_edited_profile():
            profile_data = {
                "logo_path": profile_logo_path.get(),
                "logo_height_percent": profile_logo_height_percent.get(),
                "heading_color": profile_heading_color.get(),
                "body_color": profile_body_color.get(),
                "shader_color": profile_shader_color.get(),
                "preparer_name": profile_preparer_name.get(),
                "email": profile_email.get(),
                "phone": profile_phone.get(),
                "note": profile_note.get(),
                "action": profile_action.get()
            }
            save_profile(profile_name.get(), profile_data)
            profile_window.destroy()
            load_profile_names()

        # Aligning the buttons in the same row
        Button(profile_window, text="Save Profile", font=("Arial", 14), command=save_edited_profile).grid(row=11, column=0, pady=20, padx=10, sticky='e')
        Button(profile_window, text="Cancel", font=("Arial", 14), command=profile_window.destroy).grid(row=11, column=1, pady=20, padx=10, sticky='w')

    def delete_selected_profile():
        selected_profile = profile_name_var.get()
        if messagebox.askyesno("Delete Profile", f"Are you sure you want to delete the profile '{selected_profile}'?"):
            delete_profile(selected_profile)
            load_profile_names()
            clear_profile_fields()

    def clear_profile_fields():
        title.set("")
        logo_path.set("")
        logo_height_percent.set("50")
        heading_color.set("")
        body_color.set("")
        shader_color.set("")
        preparer_name.set("")
        email.set("")
        phone.set("")
        note.set("")
        action.set("Manage")
        date_of_work.set("")

    def load_profile_names():
        profiles = load_profiles()
        profile_names = list(profiles.keys())
        profile_name_combobox['values'] = profile_names

    def select_profile_callback(event):
        select_profile(profile_name_var, {
            'title': title,
            'logo_path': logo_path,
            'logo_height_percent': logo_height_percent,
            'heading_color': heading_color,
            'body_color': body_color,
            'shader_color': shader_color,
            'preparer_name': preparer_name,
            'email': email,
            'phone': phone,
            'note': note,
            'action': action,
            'date_of_work': date_of_work
        })

    root = Tk()
    root.title("Work Log Report Generator")
    root.geometry("1200x800")

    title = StringVar()
    action = StringVar(value="Manage")
    logo_path = StringVar()
    logo_height_percent = StringVar(value="50")
    heading_color = StringVar()
    body_color = StringVar()
    shader_color = StringVar()
    preparer_name = StringVar()
    email = StringVar()
    phone = StringVar()
    note = StringVar()
    date_of_work = StringVar(value=datetime.today().strftime('%Y-%m-%d'))
    add_clipboard_image = StringVar(value="0")
    profile_name_var = StringVar()

    image_labels = [Label(root) for _ in range(5)]
    image_paths = [None] * 5
    image_entries = [Entry(root, width=30, font=("Arial", 14)) for _ in range(5)]

    Label(root, text="Title:", font=("Arial", 14)).grid(row=0, column=0, sticky='w')
    Entry(root, textvariable=title, width=50, font=("Arial", 14)).grid(row=0, column=1)

    Label(root, text="Action:", font=("Arial", 14)).grid(row=1, column=0, sticky='w')
    action_combobox = ttk.Combobox(root, textvariable=action, values=load_actions(), width=18, font=("Arial", 14))
    action_combobox.grid(row=1, column=1, sticky='w')

    Label(root, text="Logo File:", font=("Arial", 14)).grid(row=2, column=0, sticky='w')
    Entry(root, textvariable=logo_path, width=50, font=("Arial", 14)).grid(row=2, column=1)
    Button(root, text="Browse", font=("Arial", 14), command=lambda: browse_file(logo_path, [("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")])).grid(row=2, column=2)

    Label(root, text="Logo Height (%):", font=("Arial", 14)).grid(row=3, column=0, sticky='w')
    Entry(root, textvariable=logo_height_percent, width=10, font=("Arial", 14)).grid(row=3, column=1, sticky='w')

    Label(root, text="Heading Color (hex):", font=("Arial", 14)).grid(row=4, column=0, sticky='w')
    Entry(root, textvariable=heading_color, width=10, font=("Arial", 14)).grid(row=4, column=1, sticky='w')

    Label(root, text="Body Text Color (hex):", font=("Arial", 14)).grid(row=5, column=0, sticky='w')
    Entry(root, textvariable=body_color, width=10, font=("Arial", 14)).grid(row=5, column=1, sticky='w')

    Label(root, text="Column Shading Color (hex):", font=("Arial", 14)).grid(row=6, column=0, sticky='w')
    Entry(root, textvariable=shader_color, width=10, font=("Arial", 14)).grid(row=6, column=1, sticky='w')

    Label(root, text="Preparer Name:", font=("Arial", 14)).grid(row=7, column=0, sticky='w')
    Entry(root, textvariable=preparer_name, width=50, font=("Arial", 14)).grid(row=7, column=1)

    Label(root, text="Email:", font=("Arial", 14)).grid(row=8, column=0, sticky='w')
    Entry(root, textvariable=email, width=50, font=("Arial", 14)).grid(row=8, column=1)

    Label(root, text="Office Phone:", font=("Arial", 14)).grid(row=9, column=0, sticky='w')
    Entry(root, textvariable=phone, width=50, font=("Arial", 14)).grid(row=9, column=1)

    Label(root, text="Note:", font=("Arial", 14)).grid(row=10, column=0, sticky='w')
    Entry(root, textvariable=note, width=50, font=("Arial", 14)).grid(row=10, column=1)

    Label(root, text="Date of Work:", font=("Arial", 14)).grid(row=11, column=0, sticky='w')
    DateEntry(root, textvariable=date_of_work, date_pattern='yyyy-mm-dd', width=20, font=("Arial", 14)).grid(row=11, column=1, sticky='w')

    Label(root, text="Clipboard Images:", font=("Arial", 14)).grid(row=12, column=0, pady=10, sticky='nw')
    for i in range(5):
        image_entries[i].grid(row=12 + i, column=1, sticky='w')
        image_entries[i].bind("<Control-v>", lambda event, i=i: paste_image_from_clipboard(event, image_entries[i], image_labels, image_paths, i))
        Button(root, text="Clear", font=("Arial", 14), command=lambda i=i: clear_image_box(image_labels, image_paths, i)).grid(row=12 + i, column=3, sticky='w')
        image_labels[i].grid(row=12 + i, column=2, padx=5, pady=5)

    Button(root, text="Clear All Images", font=("Arial", 14), command=lambda: clear_all_images(image_labels, image_paths)).grid(row=17, column=1, pady=10, padx=10, sticky='w')
    Button(root, text="Generate Report", font=("Arial", 14), command=submit).grid(row=17, column=2, pady=20, padx=10, sticky='w')

    # Aligning the profile management buttons in the same row
    Label(root, text="Profile Management:", font=("Arial", 14)).grid(row=18, column=0, sticky='w')
    Button(root, text="Create Profile", font=("Arial", 14), command=create_profile).grid(row=18, column=1, pady=10, padx=10, sticky='w')
    Button(root, text="Edit Profile", font=("Arial", 14), command=edit_profile).grid(row=18, column=2, pady=10, padx=10, sticky='w')
    Button(root, text="Delete Profile", font=("Arial", 14), command=delete_selected_profile).grid(row=18, column=3, pady=10, padx=10, sticky='w')

    Label(root, text="Select Profile:", font=("Arial", 14)).grid(row=19, column=0, sticky='w')
    profile_name_combobox = ttk.Combobox(root, textvariable=profile_name_var, state='readonly', font=("Arial", 14))
    profile_name_combobox.grid(row=19, column=1, pady=10, padx=10)
    profile_name_combobox.bind("<<ComboboxSelected>>", select_profile_callback)

    load_profile_names()

    root.mainloop()

if __name__ == "__main__":
    main()
